#!/usr/bin/env python3
"""
Markdown -> DOCX -> PDF (+ HTML preview) pipeline for resumes, CVs, and
cover letters.

Designed to run in Cowork's code-execution sandbox (python-docx +
LibreOffice headless are pre-installed) and locally on Mac/Windows/Linux
when LibreOffice is on PATH. If `soffice` is not available, the .docx is
still written and the PDF step is skipped with a clear message — the
markdown is the canonical artifact, the docx is a useful fallback.

Each run also writes a `.html` preview next to the .docx. Open it in any
browser to eyeball the formatting without needing Word, LibreOffice, or
the PDF render step. The preview mimics page geometry (Letter, 0.5in/0.55in
margins) so what you see is roughly what the PDF will be. The docx and
PDF remain canonical — HTML is for at-a-glance review only.

Usage:
    python scripts/generate-docx.py <input.md> [<input2.md> ...]

Output: a .docx, a .html preview, and (when soffice is available) a .pdf
next to each input.
"""

import argparse
import html as html_lib
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt


# --------------------------------------------------------------------------
# Design tokens — match templates/shared.css from the prior pipeline so the
# DOCX output is visually equivalent to "single-column serif with restrained
# navy accent."
# --------------------------------------------------------------------------

NAVY = RGBColor(0x2C, 0x5F, 0x8A)
TEXT = RGBColor(0x2D, 0x2D, 0x2D)
MUTED = RGBColor(0x55, 0x55, 0x66)

# Georgia replaces Charter: widely installed across Mac/Windows/Linux,
# designed for screen reading, neutral, clearly different from Times.
# LibreOffice substitutes Liberation Serif if Georgia isn't installed,
# and embeds whichever it uses in the PDF output.
FONT = "Georgia"

BODY_SIZE_PT = 10.5
LINE_SPACING = 1.35

# Anything that survives copy-paste from web/Word but trips ATS parsers.
UNICODE_REPLACEMENTS = {
    "\u2014": "-",    # em-dash
    "\u2013": "-",    # en-dash
    "\u201C": '"',    # left double quote
    "\u201D": '"',    # right double quote
    "\u2018": "'",    # left single quote
    "\u2019": "'",    # right single quote
    "\u2026": "...",  # ellipsis
    "\u00A0": " ",    # non-breaking space
}
ZERO_WIDTH_RE = re.compile(r"[\u200B\u200C\u200D\uFEFF]")
FRONTMATTER_RE = re.compile(r"\A---\r?\n.*?\r?\n---\r?\n[\r\n]*", re.DOTALL)
HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)
ASK_VERIFY_RE = re.compile(r"\[(ASK|VERIFY):[^\]]+\]")
BRACKET_PLACEHOLDER_RE = re.compile(r"(?<!\!)\[(?![ xX]\])([^\]\n]{2,120})\](?!\()")
YEAR_TBD_RE = re.compile(r"\byear TBD\b", re.IGNORECASE)


# --------------------------------------------------------------------------
# Markdown preprocessing
# --------------------------------------------------------------------------

def normalize_unicode(text: str) -> str:
    for k, v in UNICODE_REPLACEMENTS.items():
        text = text.replace(k, v)
    return ZERO_WIDTH_RE.sub("", text)


def strip_frontmatter(text: str) -> str:
    return FRONTMATTER_RE.sub("", text, count=1)


def find_render_blockers(text: str, kind: str):
    """Return human-readable reasons the markdown should not be rendered."""
    body = strip_frontmatter(normalize_unicode(text))
    blockers = []

    if HTML_COMMENT_RE.search(body):
        blockers.append("HTML comments/template notes are still present")
        body = HTML_COMMENT_RE.sub("", body)

    if YEAR_TBD_RE.search(body):
        blockers.append("unresolved placeholder: year TBD")

    for match in ASK_VERIFY_RE.finditer(body):
        label = match.group(1).upper()
        blockers.append(f"unresolved {label} marker: {match.group(0)}")

    for match in BRACKET_PLACEHOLDER_RE.finditer(body):
        token = match.group(1).strip()
        blockers.append(f"unresolved bracket placeholder: [{token}]")

    # Resume/CV/contact markdown legitimately uses [Link Text](url). The
    # regex above excludes those, so any remaining bracket token is almost
    # certainly a template placeholder that should not ship to DOCX/PDF.
    seen = set()
    unique = []
    for blocker in blockers:
        if blocker not in seen:
            seen.add(blocker)
            unique.append(blocker)
    return unique


def validate_markdown_for_render(text: str, kind: str):
    blockers = find_render_blockers(text, kind)
    if blockers:
        joined = "; ".join(blockers)
        raise ValueError(
            f"{kind} markdown is not ready to render: {joined}"
        )


def parse_resume_sections(markdown: str):
    """Same contract as the prior JS parser:
    line 1 = '# Name', next non-empty line = contact, body starts after the
    first '---' divider.
    """
    lines = markdown.split("\n")
    h1_match = re.match(r"^#\s+(.+?)\s*$", lines[0] if lines else "")
    if not h1_match:
        raise ValueError("First line must be an h1 with the name, e.g. '# Sarah Chen'")
    name = h1_match.group(1)

    i = 1
    while i < len(lines) and lines[i].strip() == "":
        i += 1
    if i >= len(lines):
        raise ValueError("No contact line found after name")
    contact = lines[i].strip()

    divider = -1
    for j in range(i + 1, len(lines)):
        if lines[j].strip() == "---":
            divider = j
            break
    if divider == -1:
        raise ValueError("No '---' divider found after contact line")

    body = "\n".join(lines[divider + 1:]).lstrip("\n")
    return name, contact, body


def pick_kind(input_path):
    base = Path(input_path).name.lower()
    if base.startswith("coverletter"):
        return "coverletter"
    if base.startswith("cv"):
        return "cv"
    return "resume"


# --------------------------------------------------------------------------
# DOCX low-level helpers — python-docx exposes most of what we need, but
# borders, character spacing, and proper hyperlinks all require dropping to
# OOXML. These wrappers keep the OOXML noise out of the layout code.
# --------------------------------------------------------------------------

def _set_run_fonts(rPr, name=FONT):
    """Ensure all four font slots (ascii/hAnsi/eastAsia/cs) point at `name`.
    Without eastAsia, some viewers fall back to Times for any character that
    happens to land in the East-Asian Unicode range."""
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), name)


def style_run(run, *, size=None, bold=None, italic=None, color=None,
              spacing_pt=None):
    rPr = run._r.get_or_add_rPr()
    _set_run_fonts(rPr)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if spacing_pt is not None:
        # w:spacing val is in 20ths of a point.
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(round(spacing_pt * 20))))
        rPr.append(sp)


def add_bottom_border(paragraph, *, color="2C5F8A", size_eighths=4):
    """Hairline bottom border (e.g. section header underline).
    `size_eighths` is in 8ths of a point — 4 = 0.5pt, the value used in
    shared.css for both the header rule and section underlines."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_hyperlink(paragraph, text, url, *, color=None, size=None,
                  bold=False, italic=False):
    """Real DOCX hyperlink (w:hyperlink with an external relationship), not
    just a styled run. python-docx has no first-class API for this, so we
    build the OOXML directly."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    _set_run_fonts(rPr)

    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(round(size * 2))))  # half-points
        rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if color is not None:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "{:02X}{:02X}{:02X}".format(*color))
        rPr.append(col)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --------------------------------------------------------------------------
# Markdown -> DOCX walker
#
# Resume markdown is a constrained subset of CommonMark — h2/h3 sections,
# paragraphs, single-level bullets, **bold**, *italic*, [links](url), and
# the trailing-two-spaces hardbreak idiom on the skills line. We tokenize
# with markdown-it-py (which handles the parsing edge cases hand-rolled
# regex would miss — nested emphasis, escapes, code spans) and walk the
# token stream into styled python-docx paragraphs.
# --------------------------------------------------------------------------

_md = MarkdownIt("commonmark", {"breaks": False, "html": False})


def render_inline(paragraph, inline_tokens, *, default_color=None,
                  default_size=None):
    """Walk markdown-it inline tokens, emit runs into `paragraph`."""
    bold = False
    italic = False
    link_url = None

    def emit_text(text):
        if text == "":
            return
        if link_url is not None:
            add_hyperlink(
                paragraph, text, link_url,
                color=default_color, size=default_size,
                bold=bold, italic=italic,
            )
        else:
            run = paragraph.add_run(text)
            # Navy is reserved for name and section headers (set via
            # default_color by emit_header/emit_h2). Body-level **bold**
            # inherits the paragraph default so role titles and skill
            # labels stay near-black — accent in 2-3 places, never more.
            color = default_color
            style_run(run, size=default_size, bold=bold, italic=italic,
                      color=color)

    for tok in inline_tokens:
        if tok.type == "text":
            emit_text(tok.content)
        elif tok.type == "softbreak":
            emit_text(" ")
        elif tok.type == "hardbreak":
            # Markdown's "two trailing spaces" idiom — used in the skills
            # block to keep categories visually grouped.
            run = paragraph.add_run()
            style_run(run, size=default_size, color=default_color)
            run.add_break()
        elif tok.type == "strong_open":
            bold = True
        elif tok.type == "strong_close":
            bold = False
        elif tok.type == "em_open":
            italic = True
        elif tok.type == "em_close":
            italic = False
        elif tok.type == "link_open":
            link_url = dict(tok.attrs).get("href", "")
        elif tok.type == "link_close":
            link_url = None
        elif tok.type == "code_inline":
            # Resumes don't really use code spans, but tolerate them.
            emit_text(tok.content)
        # Skip image, html_inline — not used in resume markdown.


def emit_h2(doc, inline_tokens):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(13)  # 0.18in
    pf.space_after = Pt(4)    # 0.06in
    pf.keep_with_next = True
    render_inline(
        p,
        _uppercase_text_tokens(inline_tokens),
        default_color=NAVY,
        default_size=10.5,
    )
    # Tracking — letter-spacing 0.06em ≈ 0.63pt at 10.5pt body. Apply to
    # every run in the heading.
    for run in p.runs:
        rPr = run._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), "13")  # 0.65pt in 20ths of a point
        rPr.append(sp)
        run.bold = True
    add_bottom_border(p, color="2C5F8A", size_eighths=4)


def _uppercase_text_tokens(inline_tokens):
    """Section headers render uppercase. CSS does this with text-transform;
    DOCX has no direct equivalent that works reliably across viewers, so we
    uppercase the source text. Returns a shallow-copied token list."""
    out = []
    for tok in inline_tokens:
        if tok.type == "text":
            new = tok.copy() if hasattr(tok, "copy") else _shallow_copy_token(tok)
            new.content = tok.content.upper()
            out.append(new)
        else:
            out.append(tok)
    return out


def _shallow_copy_token(tok):
    from copy import copy
    return copy(tok)


def emit_h3(doc, inline_tokens):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(7)  # 0.1in
    pf.space_after = Pt(1)
    pf.keep_with_next = True
    render_inline(p, inline_tokens, default_color=TEXT, default_size=11)
    for run in p.runs:
        run.bold = True
        # h3 strong text shouldn't pick up the navy accent — override.
        run.font.color.rgb = TEXT


def emit_company_line(doc, inline_tokens):
    """The italic+dot line right after an h3: 'Jan 2022 - Present · Remote'.
    Smaller and muted."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    render_inline(p, inline_tokens, default_color=MUTED, default_size=9.75)


def emit_body_para(doc, inline_tokens):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    render_inline(p, inline_tokens)


def emit_cover_para(doc, inline_tokens):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    render_inline(p, inline_tokens)


def emit_bullet(doc, inline_tokens):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.left_indent = Inches(0.22)
    render_inline(p, inline_tokens)


def walk_tokens(doc, tokens, *, cover_letter=False):
    after_h3 = False
    i = 0
    n = len(tokens)
    while i < n:
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1])
            inline = tokens[i + 1]
            children = inline.children or []
            if level == 2:
                emit_h2(doc, children)
                after_h3 = False
            else:
                # h3 (or h1/h4+ as a fallback)
                emit_h3(doc, children)
                after_h3 = True
            i += 3  # heading_open, inline, heading_close

        elif tok.type == "paragraph_open":
            inline = tokens[i + 1]
            children = inline.children or []
            if after_h3:
                emit_company_line(doc, children)
                after_h3 = False
            elif cover_letter:
                emit_cover_para(doc, children)
            else:
                emit_body_para(doc, children)
            i += 3

        elif tok.type == "bullet_list_open":
            after_h3 = False
            i += 1
            while i < n and tokens[i].type != "bullet_list_close":
                if tokens[i].type == "list_item_open":
                    j = i + 1
                    while j < n and tokens[j].type != "list_item_close":
                        if tokens[j].type == "paragraph_open":
                            inline = tokens[j + 1]
                            emit_bullet(doc, inline.children or [])
                            j += 3
                        else:
                            j += 1
                    i = j  # at list_item_close
                i += 1
            i += 1  # consume bullet_list_close

        elif tok.type == "hr":
            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(0)
            add_bottom_border(sep, size_eighths=2)
            i += 1

        else:
            i += 1


# --------------------------------------------------------------------------
# Markdown -> HTML preview
#
# Mirrors the DOCX walker but emits a self-contained HTML file. CSS uses
# the same design tokens (Georgia, navy accent, hairlines, page geometry)
# so the preview reads as a faithful approximation of the PDF without
# requiring LibreOffice. Print stylesheet drops the page chrome so
# Ctrl+P from the browser is a reasonable secondary fallback — but
# LibreOffice's docx-faithful PDF remains canonical.
# --------------------------------------------------------------------------

HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
:root {{
  --navy: #2C5F8A;
  --text: #2D2D2D;
  --muted: #555566;
}}
* {{ box-sizing: border-box; }}
html, body {{ margin: 0; padding: 0; }}
body {{
  background: #e9eaed;
  font-family: Georgia, "Liberation Serif", "Times New Roman", serif;
  color: var(--text);
  font-size: 10.5pt;
  line-height: 1.35;
  padding: 24px 16px;
}}
.page {{
  background: #fff;
  width: 8.5in;
  min-height: 11in;
  margin: 0 auto;
  padding: 0.5in 0.55in;
  box-shadow: 0 2px 14px rgba(0,0,0,0.12);
}}
.name {{
  font-size: 22pt;
  font-weight: bold;
  color: var(--navy);
  margin: 0 0 2pt;
}}
.contact {{
  color: var(--muted);
  font-size: 10pt;
  margin: 0 0 7pt;
  padding-bottom: 2pt;
  border-bottom: 0.5pt solid var(--navy);
}}
.contact a {{ color: inherit; }}
h2 {{
  color: var(--navy);
  font-size: 10.5pt;
  font-weight: bold;
  letter-spacing: 0.06em;
  margin: 13pt 0 4pt;
  padding-bottom: 1pt;
  border-bottom: 0.5pt solid var(--navy);
  text-transform: uppercase;
}}
h3 {{
  color: var(--text);
  font-size: 11pt;
  font-weight: bold;
  margin: 7pt 0 1pt;
}}
.company-line {{
  color: var(--muted);
  font-size: 9.75pt;
  font-style: italic;
  margin: 0 0 3pt;
}}
.body-para {{ margin: 0 0 5pt; }}
.cover-para {{ margin: 0 0 10pt; }}
ul {{
  margin: 0;
  padding-left: 0.22in;
  list-style: disc outside;
}}
ul li {{ margin: 0 0 2.5pt; }}
strong {{ font-weight: bold; }}
em {{ font-style: italic; }}
a {{ color: var(--navy); }}
hr {{
  border: 0;
  border-top: 0.25pt solid var(--navy);
  margin: 6pt 0;
}}
@media print {{
  body {{ background: #fff; padding: 0; }}
  .page {{ width: auto; min-height: 0; margin: 0; padding: 0.5in 0.55in; box-shadow: none; }}
}}
</style>
</head>
<body>
<main class="page">
{body}
</main>
</body>
</html>
"""


def render_inline_html(inline_tokens):
    """Walk markdown-it inline tokens, return an HTML string."""
    parts = []
    for tok in inline_tokens:
        if tok.type == "text":
            parts.append(html_lib.escape(tok.content))
        elif tok.type == "softbreak":
            parts.append(" ")
        elif tok.type == "hardbreak":
            parts.append("<br>")
        elif tok.type == "strong_open":
            parts.append("<strong>")
        elif tok.type == "strong_close":
            parts.append("</strong>")
        elif tok.type == "em_open":
            parts.append("<em>")
        elif tok.type == "em_close":
            parts.append("</em>")
        elif tok.type == "link_open":
            href = dict(tok.attrs).get("href", "")
            parts.append(
                f'<a href="{html_lib.escape(href, quote=True)}">'
            )
        elif tok.type == "link_close":
            parts.append("</a>")
        elif tok.type == "code_inline":
            parts.append(f"<code>{html_lib.escape(tok.content)}</code>")
        # html_inline / image are intentionally skipped — resume markdown
        # doesn't use them and we don't want raw HTML passing through.
    return "".join(parts)


def walk_tokens_html(tokens, *, cover_letter=False):
    """Mirror of `walk_tokens` but emits HTML. The italic+dot line
    immediately after an h3 gets the .company-line class; subsequent
    paragraphs in cover letters get .cover-para spacing."""
    parts = []
    after_h3 = False
    i = 0
    n = len(tokens)
    while i < n:
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1])
            inline = tokens[i + 1]
            children = inline.children or []
            if level == 2:
                # Uppercase via the source-text path so emphasis tags
                # inside a heading still uppercase their text content
                # consistently (matches the DOCX behavior).
                content = render_inline_html(_uppercase_text_tokens(children))
                parts.append(f"<h2>{content}</h2>")
                after_h3 = False
            else:
                content = render_inline_html(children)
                parts.append(f"<h3>{content}</h3>")
                after_h3 = True
            i += 3

        elif tok.type == "paragraph_open":
            inline = tokens[i + 1]
            children = inline.children or []
            content = render_inline_html(children)
            if after_h3:
                cls = "company-line"
                after_h3 = False
            elif cover_letter:
                cls = "cover-para"
            else:
                cls = "body-para"
            parts.append(f'<p class="{cls}">{content}</p>')
            i += 3

        elif tok.type == "bullet_list_open":
            after_h3 = False
            parts.append("<ul>")
            i += 1
            while i < n and tokens[i].type != "bullet_list_close":
                if tokens[i].type == "list_item_open":
                    j = i + 1
                    while j < n and tokens[j].type != "list_item_close":
                        if tokens[j].type == "paragraph_open":
                            inline = tokens[j + 1]
                            children = inline.children or []
                            content = render_inline_html(children)
                            parts.append(f"<li>{content}</li>")
                            j += 3
                        else:
                            j += 1
                    i = j  # at list_item_close
                i += 1
            parts.append("</ul>")
            i += 1  # consume bullet_list_close

        elif tok.type == "hr":
            parts.append("<hr>")
            i += 1

        else:
            i += 1
    return "\n".join(parts)


def build_html(name, contact, body_md, kind):
    contact_tokens = _md.parseInline(contact)
    inline = contact_tokens[0].children if contact_tokens else []
    contact_html = render_inline_html(inline or [])
    body_tokens = _md.parse(body_md)
    body_html = walk_tokens_html(body_tokens, cover_letter=(kind == "coverletter"))
    body = "\n".join([
        f'<h1 class="name">{html_lib.escape(name)}</h1>',
        f'<p class="contact">{contact_html}</p>',
        body_html,
    ])
    return HTML_TEMPLATE.format(title=html_lib.escape(name), body=body)


# --------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------

def configure_normal_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = TEXT
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(5)
    rPr = normal.element.get_or_add_rPr()
    _set_run_fonts(rPr)


def configure_page(doc):
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def emit_header(doc, name, contact):
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name)
    style_run(name_run, size=22, bold=True, color=NAVY)

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(7)
    contact_tokens = _md.parseInline(contact)
    inline = contact_tokens[0].children if contact_tokens else []
    render_inline(contact_para, inline or [], default_color=MUTED,
                  default_size=10)
    add_bottom_border(contact_para, color="2C5F8A", size_eighths=4)


def build_document(name, contact, body_md, kind):
    doc = Document()
    configure_page(doc)
    configure_normal_style(doc)
    emit_header(doc, name, contact)
    tokens = _md.parse(body_md)
    walk_tokens(doc, tokens, cover_letter=(kind == "coverletter"))
    return doc


# --------------------------------------------------------------------------
# Pipeline
# --------------------------------------------------------------------------

def find_soffice():
    """Return path to a usable LibreOffice CLI, or None if unavailable.
    Checks PATH first, then the standard install location on Windows."""
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    candidates = [
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    return None


def md_to_docx(input_path):
    raw = Path(input_path).read_text(encoding="utf-8")
    kind = pick_kind(input_path)
    validate_markdown_for_render(raw, kind)
    raw = strip_frontmatter(raw)
    raw = normalize_unicode(raw)
    name, contact, body = parse_resume_sections(raw)
    doc = build_document(name, contact, body, kind)
    out_path = Path(input_path).with_suffix(".docx")
    doc.save(str(out_path))
    return out_path


def md_to_html(input_path):
    raw = Path(input_path).read_text(encoding="utf-8")
    kind = pick_kind(input_path)
    validate_markdown_for_render(raw, kind)
    raw = strip_frontmatter(raw)
    raw = normalize_unicode(raw)
    name, contact, body = parse_resume_sections(raw)
    html_str = build_html(name, contact, body, kind)
    out_path = Path(input_path).with_suffix(".html")
    out_path.write_text(html_str, encoding="utf-8")
    return out_path


def docx_to_pdf_batch(docx_paths, soffice_path):
    """Convert a batch of DOCX files to PDF in a single LibreOffice process.
    Returns a list of (docx_path, pdf_path, ok, error) tuples in the same
    order as `docx_paths`."""
    if not docx_paths:
        return []
    # soffice writes PDFs to --outdir. We group by output directory so each
    # file lands next to its source. In practice every input shares the
    # same parent (we're called per-batch from main()), so this is usually
    # one process; we still split if directories differ.
    by_dir = {}
    for p in docx_paths:
        by_dir.setdefault(p.parent, []).append(p)

    results = []
    for outdir, paths in by_dir.items():
        cmd = [
            soffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(outdir),
            *[str(p) for p in paths],
        ]
        try:
            proc = subprocess.run(
                cmd, capture_output=True, text=True, timeout=120,
            )
        except subprocess.TimeoutExpired as e:
            for p in paths:
                results.append((p, p.with_suffix(".pdf"), False,
                                f"soffice timed out after 120s: {e}"))
            continue
        except OSError as e:
            for p in paths:
                results.append((p, p.with_suffix(".pdf"), False,
                                f"soffice failed to launch: {e}"))
            continue

        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or "").strip()
            for p in paths:
                pdf = p.with_suffix(".pdf")
                results.append((p, pdf, pdf.exists(),
                                None if pdf.exists() else err))
            continue

        for p in paths:
            pdf = p.with_suffix(".pdf")
            if pdf.exists():
                results.append((p, pdf, True, None))
            else:
                results.append((p, pdf, False,
                                "soffice exited 0 but no PDF produced"))
    return results


def main(argv=None):
    parser = argparse.ArgumentParser(
        prog="generate-docx.py",
        description="Markdown -> DOCX -> PDF for resumes, CVs, and cover letters",
    )
    parser.add_argument("inputs", nargs="+", help="One or more markdown files")
    args = parser.parse_args(argv)

    docx_results = []
    any_failed = False

    for raw_input in args.inputs:
        input_path = Path(raw_input).resolve()
        if not input_path.exists():
            print(f"Failed to prepare {input_path}: file not found",
                  file=sys.stderr)
            any_failed = True
            continue
        try:
            docx_path = md_to_docx(input_path)
            print(f"Wrote {docx_path}")
            docx_results.append(docx_path)
        except Exception as e:
            print(
                f"Failed to build .docx for {input_path}: {e}. "
                "If this is a markdown-not-ready error, fix the markdown and rerun.",
                  file=sys.stderr)
            any_failed = True
            continue

        # HTML preview is best-effort — a failure here does not block the
        # canonical .docx/.pdf artifacts.
        try:
            html_path = md_to_html(input_path)
            print(f"Wrote {html_path}")
        except Exception as e:
            print(
                f"HTML preview failed for {input_path}: {e} "
                "(docx/PDF unaffected)",
                file=sys.stderr,
            )

    soffice = find_soffice()
    if not docx_results:
        return 1 if any_failed else 0

    if soffice is None:
        print(
            "\nLibreOffice not found — skipping PDF conversion.\n"
            "  The .docx files above are valid and submittable. To enable\n"
            "  automatic PDF generation, install LibreOffice and ensure\n"
            "  `soffice` is on PATH:\n"
            "    macOS:   brew install --cask libreoffice\n"
            "    Windows: winget install TheDocumentFoundation.LibreOffice\n"
            "    Linux:   apt install libreoffice  (or your package manager)",
            file=sys.stderr,
        )
        return 1 if any_failed else 0

    pdf_results = docx_to_pdf_batch(docx_results, soffice)
    for docx_path, pdf_path, ok, err in pdf_results:
        if ok:
            size = pdf_path.stat().st_size
            print(f"Wrote {pdf_path} ({size} bytes)")
        else:
            any_failed = True
            print(f"PDF conversion failed for {pdf_path}: {err}",
                  file=sys.stderr)

    return 1 if any_failed else 0


if __name__ == "__main__":
    sys.exit(main())
